from django.shortcuts import render
from django.views import generic
from .models import *
from django.urls import reverse_lazy
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from .uploads import handle_uploaded_tests, handle_uploaded_properties
from .forms import *
from itertools import chain
from django.views.generic import ListView
from django.shortcuts import redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from .filters import ReportsFilter
from django_filters.views import FilterView
from django.http import HttpResponse

from docx import Document
import pypandoc

import io

import csv


@login_required(login_url='/accounts/login/')
def index(request):
    """View function for home page of site."""

    # Generate counts of some of the main objects
    num_visits = request.session.get('num_visits', 0)
    num_report = Report.objects.all().count()
    request.session['num_visits'] = num_visits + 1

    context = {
        'num_report': num_report,
        'num_visits': num_visits,
    }

    # Render the HTML template index.html with the data in the context variable
    return render(request, 'index.html', context=context)


# 1. Организации
class ReportTypeListView(LoginRequiredMixin, generic.ListView):
    model = ReportType


class ReportTypeDetailView(LoginRequiredMixin, generic.DetailView):
    model = ReportType


class ReportTypeCreate(LoginRequiredMixin, CreateView):
    model = ReportType
    form_class = ReportTypeForm


class ReportTypeUpdate(LoginRequiredMixin, UpdateView):
    model = ReportType
    form_class = ReportTypeForm


class ReportTypeDelete(LoginRequiredMixin, DeleteView):
    model = ReportType
    success_url = reverse_lazy('report_types')


# 2. Проект
class SourceListView(LoginRequiredMixin, generic.ListView):
    model = Source


class SourceDetailView(LoginRequiredMixin, generic.DetailView):
    model = Source


class SourceCreate(LoginRequiredMixin, CreateView):
    model = Source
    form_class = SourceForm


class SourceUpdate(LoginRequiredMixin, UpdateView):
    model = Source
    form_class = SourceForm


class SourceDelete(LoginRequiredMixin, DeleteView):
    model = Source
    success_url = reverse_lazy('sources')


# 3. Местрождение
class RegionListView(LoginRequiredMixin, generic.ListView):
    model = Region


class RegionDetailView(LoginRequiredMixin, generic.DetailView):
    model = Region


class RegionCreate(LoginRequiredMixin, CreateView):
    model = Region
    form_class = RegionForm


class RegionUpdate(LoginRequiredMixin, UpdateView):
    model = Region
    form_class = RegionForm


class RegionDelete(LoginRequiredMixin, DeleteView):
    model = Region
    success_url = reverse_lazy('regions')


# 4. Скважины
class ReportListView(LoginRequiredMixin, generic.ListView):
    model = Report


class ReportDetailView(LoginRequiredMixin, generic.DetailView):
    model = Report


class ReportCreate(LoginRequiredMixin, CreateView):
    model = Report
    form_class = ReportForm


class ReportUpdate(LoginRequiredMixin, UpdateView):
    model = Report
    form_class = ReportForm


class ReportDelete(LoginRequiredMixin, DeleteView):
    model = Report
    success_url = reverse_lazy('reports')


class SearchView(LoginRequiredMixin, ListView):
    template_name = 'search.html'
    paginate_by = 20
    count = 0

    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        context['count'] = self.count or 0
        context['query'] = self.request.GET.get('q')
        return context

    def get_queryset(self):
        request = self.request
        query = request.GET.get('q', None)

        if query is not None:
            specimen_results = Specimen.objects.search(query)
            property_results = Property.objects.search(query)

            # combine querysets
            queryset_chain = chain(
                specimen_results,
                property_results
            )
            qs = sorted(queryset_chain,
                        key=lambda instance: instance.pk,
                        reverse=True)
            self.count = len(qs)  # since qs is actually a list
            return qs
        return Property.objects.none()  # just an empty queryset as default


def generate_docx_file(object_id):

    # Fetch your data from the database
    report = Report.objects.get(pk=object_id)

    # Create a Word document
    document = Document()
    document.add_heading(report.report_name, 0)

    # Add database data to the document
    document.add_heading(report.chapter_name_1, 1)
    document.add_paragraph(report.chapter_text_1)

    document.add_heading(report.chapter_name_2, 1)
    document.add_paragraph(report.chapter_text_2)

    document.add_heading(report.chapter_name_3, 1)
    document.add_paragraph(report.chapter_text_3)

    document.add_heading(report.chapter_name_4, 1)
    document.add_paragraph(report.chapter_text_4)

    document.add_heading(report.chapter_name_5, 1)
    document.add_paragraph(report.chapter_text_5)

    document.add_heading(report.chapter_name_6, 1)
    document.add_paragraph(report.chapter_text_6)
    # Add more item details as needed
    return document


def export_to_docx(request, object_id):

    report = Report.objects.get(pk=object_id)

    document = generate_docx_file(object_id)
    # Save the document to an in-memory stream
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Return the document as an HTTP response
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="report{report.id}.docx"'
    return response

def export_to_pdf(request, object_id):


    report = Report.objects.get(pk=object_id)

    document = generate_docx_file(object_id)
    # Save the document to an in-memory stream
    temp_doc_path = 'documents/temp_document.docx'
    temp_pdf_path = 'documents/temp_document.pdf'

    document.save(temp_doc_path)
    
    # Convert the docx file to PDF
    output_pdf = pypandoc.convert_file(
        source_file=temp_doc_path,
        format='docx',
        to='pdf',
        outputfile=temp_pdf_path,
        extra_args=['--pdf-engine=xelatex', '-V','mainfont="Arial"','-V','geometry:margin=1in']
    )

    with open(temp_pdf_path, 'rb') as f:
        pdf_content = f.read()

    # Return the PDF as an HTTP response
    response = HttpResponse(pdf_content, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="report{report.id}.pdf"'
    return response